import traceback
import io

import base64
from typing import List

import docx
from docx.shared import Inches, Pt


import markdown

        
def blue(run):
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)

def red(run):
    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

def convert(doc:List[dict]) -> bytes:
    renderer = docx_renderer()
    renderer.digest(doc)
    return renderer.doc_to_bytes()

class docx_renderer():
    def __init__(self, template_path:str=None, make_blue=False) -> None:
        self.d = docx.Document(template_path)
        self.make_blue = make_blue

    def add_paragraph(self, newtext, *args, **kwargs):
        new_paragraph = self.d.add_paragraph(newtext, *args, **kwargs)
        if self.make_blue:
            for r in new_paragraph.runs:
                blue(r)
        return new_paragraph
    
    def add_run(self, text, *args, **kwargs):
        if not self.d.paragraphs:
            self.add_paragraph('')

        last_paragraph = self.d.paragraphs[-1]
        
        if not last_paragraph.runs:
            last_run = last_paragraph.add_run(text)
        else:
            last_run = last_paragraph.runs[-1]
            last_run.add_text(text)
            
        if self.make_blue:
            blue(last_run)
        return last_run
        
    def digest_text(self, children, *args, **kwargs):
        return self.add_paragraph(children)
    

    def digest_str(self, children, *args, **kwargs):
        return self.add_run(children)

    def digest_markdown(self, children, *args, **kwargs):
        return self.add_paragraph(children, style='Normal')
        
    def digest_verbatim(self, children, *args, **kwargs):
        new_run = self.add_run(children)
        new_run.font.name = 'Courier New'  # Or any other monospace font
        new_run.font.size = docx.shared.Pt(8)  # Adjust font size as needed
        return new_run


    def digest_error(self, children, *args, **kwargs):
        if isinstance(children, BaseException):
            traceback.print_exc(limit=5)
            children = traceback.format_exc(limit=5)

        new_run = self.add_run(children)
        new_run.font.name = 'Courier New'  # Or any other monospace font
        new_run.font.size = docx.shared.Pt(8)  # Adjust font size as needed
        red(new_run)
        return new_run


    def digest_iterator(self, children, *args, **kwargs):
        if children:
            return [self.digest(val, *args, **kwargs) for val in children]
        return []


    def digest_image(self, children, *args, **kwargs):

        image_width = Inches(max(1, kwargs.get('width', 0.8)*5))
        image_caption = kwargs.get('caption', '')
        image_blob = kwargs.get('imageblob', '')

        assert image_blob, 'no image data given!'

        # Decode the base64 image
        img_bytes = base64.b64decode(image_blob)

        # Create an image stream from the bytes
        image_stream = io.BytesIO(img_bytes)
        
        picture = self.d.add_picture(image_stream, width=image_width)
        # picture.width = image_width  # Ensure fixed width
        # picture.height = None  # Adjust height automatically
        picture.alignment = 1

        run = self.add_paragraph(image_caption)
        # run.style = 'Caption'  # Apply the 'Caption' style for formatting

        return run

    def digest(self, children, *args, **kwargs):
        try:
            # print(f'{type(children)=}, {args=} {kwargs=}')
            if not children:
                return ''
            elif isinstance(children, str):
                ret = self.digest_str(children, *args, **kwargs)
            elif isinstance(children, dict) and 'typ' in children and children['typ'] == 'iter':
                ret = self.digest_iterator(children, *args, **kwargs)
            elif isinstance(children, list) and children:
                ret = self.digest_iterator(children, *args, **kwargs)
            elif isinstance(children, dict) and 'typ' in children and children['typ'] == 'image':
                ret = self.digest_image(*args, **kwargs, **children)
            elif isinstance(children, dict) and 'typ' in children and children['typ'] == 'text':
                ret = self.digest_text(*args, **kwargs, **children)
            elif isinstance(children, dict) and 'typ' in children and children['typ'] == 'verbatim':
                ret = self.digest_verbatim(*args, **kwargs, **children)
            elif isinstance(children, dict) and 'typ' in children and children['typ'] == 'markdown':
                ret = self.digest_markdown(*args, **kwargs, **children)
            else:
                val = f'the element of type {type(children)}, could not be parsed.'
                ret = self.digest_error(val, *args, **kwargs)
        except Exception as err:
            ret = self.digest_error(err, *args, **kwargs)

        return ret

    def doc_to_bytes(self):
        with io.BytesIO() as fp:
            self.d.save(fp)
            fp.seek(0)
            return fp.read()

    def save(self, filepath):
        self.d.save(filepath)
